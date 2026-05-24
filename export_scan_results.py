"""
export_scan_results.py
----------------------
Escanea las fotos 2.jpeg, 3.jpeg y 4.jpeg con OCR y exporta
los datos parseados a: Excel, Word (.docx), TXT y PDF.
Los archivos se guardan en la carpeta exports/.

Uso:  py export_scan_results.py
"""

import os
import sys

# ── asegurar que los imports del proyecto funcionen ──────────────────────────
ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, ROOT)

from utils.ocr_scan import parse_department_report_image

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                 Paragraph, Spacer, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ── configuración ─────────────────────────────────────────────────────────────
PHOTOS = ["2.jpeg", "3.jpeg", "4.jpeg"]
OUT_DIR = os.path.join(ROOT, "exports")
os.makedirs(OUT_DIR, exist_ok=True)

# Cabeceras según tipo de reporte
PLU_HEADERS  = ["Descripción", "Depto.", "Count", "Price", "Sales", "% Dept", "% Total"]
DEPT_HEADERS = ["Departamento", "Ventas Netas", "Impuesto", "Total", "% Total"]

# ─────────────────────────────────────────────────────────────────────────────
# Utilidades de extracción de fila
# ─────────────────────────────────────────────────────────────────────────────

def row_to_plu(r):
    """Convierte un dict de fila PLU en lista de celdas para la tabla."""
    return [
        r.get("description", ""),
        r.get("dept_num", ""),
        r.get("items", 0),
        f"${r.get('sales_gross', 0.0):.2f}",
        f"${r.get('refunds', 0.0):.2f}",
        f"{r.get('discounts', 0.0):.2f}%",
        f"{r.get('net_sales', 0.0):.2f}%",
    ]

def row_to_dept(r):
    """Convierte un dict de fila DEPT en lista de celdas."""
    return [
        r.get("department", r.get("description", "")),
        f"${r.get('net_sales', r.get('sales', 0.0)):.2f}",
        f"${r.get('tax', 0.0):.2f}",
        f"${r.get('total', r.get('sales_gross', 0.0)):.2f}",
        f"{r.get('pct_total', r.get('net_sales', 0.0)):.2f}%",
    ]

def get_headers(report_type):
    return PLU_HEADERS if report_type == "plu" else DEPT_HEADERS

def get_cells(r, report_type):
    return row_to_plu(r) if report_type == "plu" else row_to_dept(r)

# ─────────────────────────────────────────────────────────────────────────────
# PASO 1 — OCR + parsing
# ─────────────────────────────────────────────────────────────────────────────

print("=" * 60)
print("  ESCANEANDO FOTOS CON OCR...")
print("=" * 60)

results = []   # lista de dicts: {photo, report_type, rows, parse_errors}

for photo in PHOTOS:
    path = os.path.join(ROOT, photo)
    if not os.path.exists(path):
        print(f"  [OMITIDA] {photo} — archivo no encontrado")
        continue
    print(f"\n  Procesando {photo} ...", end=" ", flush=True)
    try:
        data = parse_department_report_image(path)
        results.append({
            "photo":        photo,
            "report_type":  data.get("report_type", "plu"),
            "rows":         data.get("rows", []),
            "parse_errors": data.get("parse_errors", []),
            "raw_text":     data.get("raw_text", ""),
        })
        print(f"OK — {len(data.get('rows',[]))} filas, tipo={data.get('report_type','?')}")
    except Exception as e:
        print(f"ERROR: {e}")
        results.append({
            "photo": photo, "report_type": "?",
            "rows": [], "parse_errors": [str(e)], "raw_text": "",
        })

# ─────────────────────────────────────────────────────────────────────────────
# PASO 2 — Exportar TXT
# ─────────────────────────────────────────────────────────────────────────────

txt_path = os.path.join(OUT_DIR, "scan_results.txt")
with open(txt_path, "w", encoding="utf-8") as f:
    f.write("RESULTADOS DEL ESCÁNER OCR — MiGas App\n")
    f.write("=" * 80 + "\n\n")

    for res in results:
        f.write(f"FOTO: {res['photo']}  |  Tipo: {res['report_type'].upper()}\n")
        f.write("-" * 80 + "\n")
        headers = get_headers(res["report_type"])
        col_w = [28, 14, 6, 10, 10, 9, 9]  # anchos para PLU; DEPT usa los primeros
        col_w = col_w[:len(headers)]

        # Encabezado
        f.write("".join(h.ljust(col_w[i]) for i, h in enumerate(headers)) + "\n")
        f.write("".join("-" * w for w in col_w) + "\n")

        if not res["rows"]:
            f.write("  (sin filas)\n")
        for r in res["rows"]:
            cells = get_cells(r, res["report_type"])
            line = "".join(str(c).ljust(col_w[i]) for i, c in enumerate(cells))
            f.write(line + "\n")

        if res["parse_errors"]:
            f.write("\n  ERRORES DE PARSING:\n")
            for e in res["parse_errors"]:
                f.write(f"    • {e}\n")
        f.write("\n")

print(f"\n  [TXT]   {txt_path}")

# ─────────────────────────────────────────────────────────────────────────────
# PASO 3 — Exportar Excel (.xlsx)
# ─────────────────────────────────────────────────────────────────────────────

HEADER_FILL  = PatternFill("solid", fgColor="1F4E79")
HEADER_FONT  = Font(bold=True, color="FFFFFF", size=10)
TITLE_FONT   = Font(bold=True, size=12, color="1F4E79")
ALT_FILL     = PatternFill("solid", fgColor="D9E1F2")
BORDER_SIDE  = Side(style="thin", color="BFBFBF")
CELL_BORDER  = Border(left=BORDER_SIDE, right=BORDER_SIDE,
                      top=BORDER_SIDE, bottom=BORDER_SIDE)

wb = openpyxl.Workbook()
wb.remove(wb.active)   # quitar hoja vacía por defecto

for res in results:
    sheet_name = res["photo"].replace(".", "_")
    ws = wb.create_sheet(title=sheet_name)

    # Título
    ws.append([f"Foto: {res['photo']}  —  Reporte: {res['report_type'].upper()}"])
    title_cell = ws.cell(row=1, column=1)
    title_cell.font = TITLE_FONT
    headers = get_headers(res["report_type"])
    ws.merge_cells(start_row=1, start_column=1,
                   end_row=1, end_column=len(headers))
    title_cell.alignment = Alignment(horizontal="center")
    ws.append([])   # fila vacía

    # Cabeceras
    header_row = 3
    for col_idx, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=h)
        cell.font   = HEADER_FONT
        cell.fill   = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")
        cell.border = CELL_BORDER

    # Datos
    for row_idx, r in enumerate(res["rows"], start=1):
        cells = get_cells(r, res["report_type"])
        excel_row = header_row + row_idx
        fill = ALT_FILL if row_idx % 2 == 0 else None
        for col_idx, val in enumerate(cells, start=1):
            cell = ws.cell(row=excel_row, column=col_idx, value=val)
            cell.border = CELL_BORDER
            cell.alignment = Alignment(horizontal="left" if col_idx <= 2 else "center")
            if fill:
                cell.fill = fill

    if not res["rows"]:
        ws.cell(row=header_row + 1, column=1, value="(sin filas)")

    # Errores
    if res["parse_errors"]:
        err_row = header_row + len(res["rows"]) + 2
        ws.cell(row=err_row, column=1,
                value="Errores de parsing:").font = Font(bold=True, color="C00000")
        for i, e in enumerate(res["parse_errors"]):
            ws.cell(row=err_row + 1 + i, column=1, value=f"• {e}")

    # Ajustar ancho de columnas
    for col_idx in range(1, len(headers) + 1):
        max_len = len(headers[col_idx - 1])
        for row in ws.iter_rows(min_row=header_row, max_col=col_idx, min_col=col_idx):
            for cell in row:
                try:
                    max_len = max(max_len, len(str(cell.value or "")))
                except Exception:
                    pass
        ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)

xlsx_path = os.path.join(OUT_DIR, "scan_results.xlsx")
wb.save(xlsx_path)
print(f"  [XLSX]  {xlsx_path}")

# ─────────────────────────────────────────────────────────────────────────────
# PASO 4 — Exportar Word (.docx)
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Estilos de página — orientación apaisada para que quepan las columnas
section = doc.sections[0]
section.page_width  = Inches(11)
section.page_height = Inches(8.5)
section.left_margin = section.right_margin = Inches(0.6)
section.top_margin  = section.bottom_margin = Inches(0.5)

# Título general
title_para = doc.add_heading("Resultados del Escáner OCR — MiGas App", level=1)
title_para.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

for res in results:
    doc.add_heading(f"Foto: {res['photo']}  |  Tipo: {res['report_type'].upper()}", level=2)

    if not res["rows"]:
        doc.add_paragraph("(sin filas)")
    else:
        headers = get_headers(res["report_type"])
        tbl = doc.add_table(rows=1 + len(res["rows"]), cols=len(headers))
        tbl.style = "Table Grid"

        # Fila de encabezados
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size  = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Fondo azul en cabecera — inyectado a nivel XML
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1F4E79")
            tc_pr.append(shd)

        # Filas de datos
        for row_idx, r in enumerate(res["rows"], start=1):
            cells_data = get_cells(r, res["report_type"])
            row_cells = tbl.rows[row_idx].cells
            bg = "D9E1F2" if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, val in enumerate(cells_data):
                cell = row_cells[col_idx]
                cell.text = str(val)
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(8)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
                )
                # Fondo alterno
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), bg)
                tc_pr.append(shd)

    if res["parse_errors"]:
        p = doc.add_paragraph()
        run = p.add_run("Errores de parsing: " + " | ".join(res["parse_errors"]))
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run.font.size = Pt(8)

    doc.add_paragraph()   # espacio entre fotos

docx_path = os.path.join(OUT_DIR, "scan_results.docx")
doc.save(docx_path)
print(f"  [DOCX]  {docx_path}")

# ─────────────────────────────────────────────────────────────────────────────
# PASO 5 — Exportar PDF (reportlab)
# ─────────────────────────────────────────────────────────────────────────────

pdf_path = os.path.join(OUT_DIR, "scan_results.pdf")

doc_pdf = SimpleDocTemplate(
    pdf_path,
    pagesize=landscape(letter),
    leftMargin=0.5 * inch, rightMargin=0.5 * inch,
    topMargin=0.5 * inch, bottomMargin=0.5 * inch,
)

styles = getSampleStyleSheet()
title_style = ParagraphStyle("TitleCustom", parent=styles["Title"],
                              textColor=colors.HexColor("#1F4E79"),
                              fontSize=14, spaceAfter=6)
h2_style = ParagraphStyle("H2Custom", parent=styles["Heading2"],
                           textColor=colors.HexColor("#1F4E79"),
                           fontSize=11, spaceBefore=12, spaceAfter=4)
err_style = ParagraphStyle("ErrCustom", parent=styles["Normal"],
                            textColor=colors.red, fontSize=7)

HDR_BG  = colors.HexColor("#1F4E79")
ALT_BG  = colors.HexColor("#D9E1F2")
WHT_BG  = colors.white
HDR_TXT = colors.white
DAT_TXT = colors.black

story = []
story.append(Paragraph("Resultados del Escáner OCR — MiGas App", title_style))
story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#1F4E79")))
story.append(Spacer(1, 0.1 * inch))

for res in results:
    story.append(Paragraph(
        f"Foto: {res['photo']}  |  Tipo: {res['report_type'].upper()}  |  "
        f"Filas: {len(res['rows'])}",
        h2_style,
    ))

    headers = get_headers(res["report_type"])

    if not res["rows"]:
        story.append(Paragraph("(sin filas)", styles["Normal"]))
    else:
        table_data = [headers]
        for r in res["rows"]:
            table_data.append([str(c) for c in get_cells(r, res["report_type"])])

        # Anchos de columna proporcionales a la página (10 pulgadas útiles)
        if res["report_type"] == "plu":
            col_widths = [2.4*inch, 1.4*inch, 0.5*inch, 0.9*inch,
                          0.9*inch, 0.8*inch, 0.8*inch]
        else:
            col_widths = [2.5*inch, 1.5*inch, 1.2*inch, 1.2*inch, 1.0*inch]

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)

        ts = TableStyle([
            # Cabecera
            ("BACKGROUND",  (0, 0), (-1, 0),  HDR_BG),
            ("TEXTCOLOR",   (0, 0), (-1, 0),  HDR_TXT),
            ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, 0),  8),
            ("ALIGN",       (0, 0), (-1, 0),  "CENTER"),
            ("VALIGN",      (0, 0), (-1, -1), "MIDDLE"),
            # Datos
            ("FONTNAME",    (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",    (0, 1), (-1, -1), 7.5),
            ("ALIGN",       (0, 1), (1, -1),  "LEFT"),
            ("ALIGN",       (2, 1), (-1, -1), "CENTER"),
            # Rejilla
            ("GRID",        (0, 0), (-1, -1), 0.4, colors.HexColor("#BFBFBF")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHT_BG, ALT_BG]),
            ("TOPPADDING",  (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ])
        tbl.setStyle(ts)
        story.append(tbl)

    if res["parse_errors"]:
        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph(
            "Errores: " + " | ".join(res["parse_errors"]), err_style,
        ))

    story.append(Spacer(1, 0.15 * inch))

doc_pdf.build(story)
print(f"  [PDF]   {pdf_path}")

# ─────────────────────────────────────────────────────────────────────────────
print("\n" + "=" * 60)
print("  EXPORTACIÓN COMPLETA")
print(f"  Carpeta: {OUT_DIR}")
print("=" * 60)
